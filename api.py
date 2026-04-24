from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from src.graph import build_graph
from docx import Document
import tempfile
import os

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_methods=["*"],
    allow_headers=["*"],
)

class AnalysisRequest(BaseModel):
    cv_texto: str
    oferta_texto: str

class DownloadRequest(BaseModel):
    cv_adaptado: str

@app.post("/analyze")
async def analyze(request: AnalysisRequest):
    graph = build_graph()
    result = graph.invoke({
        "cv_texto": request.cv_texto,
        "oferta_texto": request.oferta_texto,
        "skills_cv": [],
        "skills_oferta": [],
        "skills_match": [],
        "skills_gap": [],
        "fit_score": 0.0,
        "recomendaciones": "",
        "cv_adaptado": ""
    })
    return {
        "fit_score": result["fit_score"],
        "skills_match": result["skills_match"],
        "skills_gap": result["skills_gap"],
        "recomendaciones": result["recomendaciones"],
        "cv_adaptado": result["cv_adaptado"]
    }

@app.post("/download")
async def download(request: DownloadRequest):
    doc = Document()
    for line in request.cv_adaptado.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=1)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=0)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.replace("**", "")).bold = True
        else:
            doc.add_paragraph(line)
    
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return FileResponse(tmp.name, filename="cv_adaptado.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")